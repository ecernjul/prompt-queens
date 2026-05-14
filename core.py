"""
Prompt Queens Content Generator — Core Module

Shared logic used by both the CLI (generate_content.py) and
the Streamlit UI (app.py).
"""

import ast
import io
import os
import re
from datetime import datetime
from pathlib import Path

os.environ["TRANSFORMERS_VERBOSITY"] = "error"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import anthropic
import torch
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from pinecone import Pinecone
from transformers import CLIPModel, CLIPProcessor

# ── Constants ─────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
ENV_FILE = BASE_DIR / ".env"
BRAND_VOICE_FILE = BASE_DIR / "brand_voice.txt"
INDEX_NAME = "salsify-products-clip"
CLIP_MODEL_NAME = "openai/clip-vit-base-patch32"
CLAUDE_MODEL = "claude-sonnet-4-6"

SECTION_KEYS = [
    "PDP Copy",
    "SEO Keywords",
    "Organic Social Post",
    "Outbound Email",
    "Headlines",
    "Paid Social Ad",
    "Video Script",
    "Creative Brief",
    "Competitive SWOT",
]

# ── Environment ───────────────────────────────────────────────────────────────

def load_env():
    if not ENV_FILE.exists():
        return
    with open(ENV_FILE) as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip())


def load_brand_voice() -> str:
    if not BRAND_VOICE_FILE.exists():
        return ""
    lines = []
    with open(BRAND_VOICE_FILE) as f:
        for line in f:
            stripped = line.rstrip()
            if stripped.startswith("#"):
                continue
            lines.append(stripped)
    text = "\n".join(lines).strip()
    if not text or "[PASTE EXAMPLE COPY HERE]" in text:
        return ""
    return text


# ── Model / client loaders ────────────────────────────────────────────────────

def load_clip_model():
    """Load CLIP model and processor. Call once and cache the result."""
    model = CLIPModel.from_pretrained(CLIP_MODEL_NAME)
    processor = CLIPProcessor.from_pretrained(CLIP_MODEL_NAME)
    model.eval()
    return model, processor


def get_pinecone_index():
    load_env()
    api_key = os.environ.get("PINECONE_API_KEY", "")
    if not api_key:
        raise ValueError("PINECONE_API_KEY is not set. Add it to Railway's environment variables.")
    pc = Pinecone(api_key=api_key)
    return pc.Index(INDEX_NAME)


def get_anthropic_client():
    load_env()
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY is not set. Add it to Railway's environment variables.")
    return anthropic.Anthropic(api_key=api_key)


# ── Product helpers ───────────────────────────────────────────────────────────

def parse_localized(value) -> str:
    if not value:
        return ""
    s = str(value)
    if s.startswith("{") and "en-US" in s:
        try:
            d = ast.literal_eval(s)
            if isinstance(d, dict):
                return d.get("en-US", s)
        except Exception:
            pass
    return s


def format_product(meta: dict) -> str:
    fields = {
        "SKU": meta.get("Product_Code", ""),
        "Name": parse_localized(meta.get("Product_Description", "")),
        "Short Description": parse_localized(meta.get("Short_Description_50", "")),
        "Full Description": parse_localized(meta.get("Alternate_Product_Description", "")),
        "Brand Agnostic Description": parse_localized(meta.get("Brand_Agnostic_Product_Description", "")),
        "Features": parse_localized(meta.get("Features", "")),
        "Color": parse_localized(meta.get("Color", "")),
        "Material": parse_localized(meta.get("Material", "")),
        "Weight Capacity": parse_localized(meta.get("Weight_Capacity", "")),
        "Dimensions": parse_localized(meta.get("Dimensions", "")),
        "Assembly Required": parse_localized(meta.get("Assembly_Required", "")),
        "Brand": parse_localized(meta.get("Brand", "")),
        "Category": parse_localized(meta.get("Category", "")),
    }
    return "\n".join(f"{k}: {v.strip()}" for k, v in fields.items() if v and v.strip())


def product_display_name(meta: dict, fallback: str = "") -> str:
    return parse_localized(
        meta.get("Product_Description",
        meta.get("Short_Description_50", fallback))
    )


# ── Search ────────────────────────────────────────────────────────────────────

def embed_text(text: str, model, processor) -> list[float]:
    inputs = processor(text=[text], return_tensors="pt", padding=True, truncation=True)
    with torch.no_grad():
        outputs = model.text_model(**{k: v for k, v in inputs.items() if k != "pixel_values"})
        emb = model.text_projection(outputs.pooler_output)
    emb = emb / emb.norm(dim=-1, keepdim=True)
    return emb.squeeze().tolist()


def search_products(query: str, sku_mode: bool, top_k: int, index, clip_model, clip_processor) -> list[dict]:
    embedding = embed_text(query, clip_model, clip_processor)
    if sku_mode:
        # Exact match only — no fallback to text search
        results = index.query(
            vector=embedding,
            top_k=1,
            include_metadata=True,
            filter={"Product_Code": {"$eq": query}},
        )
    else:
        results = index.query(vector=embedding, top_k=top_k, include_metadata=True)
    return results["matches"]


# ── Claude ────────────────────────────────────────────────────────────────────

def build_system_prompt(brand_voice: str) -> str:
    voice_block = ""
    if brand_voice:
        voice_block = f"""
BRAND VOICE EXAMPLES
Study these carefully and match their tone, sentence length, word choice, and energy:

{brand_voice}

---
"""
    return f"""You are an expert e-commerce copywriter for Flash Furniture / The Ubique Group — a furniture and home goods brand.
{voice_block}
Your job is to generate nine types of marketing content from product data. Follow these style guidelines:

- PDP Copy: benefit-led, scannable, 150–250 words. Lead → bullet features → "who it's for."
- SEO Keywords: 15–20 terms, purchase-intent focused, mix of head (2-3 words) and long-tail (4-6 words).
- Organic Social Post: warm, conversational, under 150 words, 5-8 hashtags. Not corporate.
- Outbound Email: B2B tone targeting designers, hospitality buyers, or facility managers. Include Subject line. Under 200 words.
- Headlines: 8 variations for ads, email subjects, banners. Under 10 words each.
- Paid Social Ad: hook + 2-3 benefits + clear CTA. Under 100 words. 3-5 hashtags. More direct than organic post.
- Video Script: 30-45 second script. Use [VISUAL:] and [VO:] cues. Punchy and visual.
- Creative Brief: 3 photography/design concepts. Format each header as **Concept 1: Title**, **Concept 2: Title**, **Concept 3: Title** (bold, on its own line). Follow each with 2-3 sentences covering setting, mood, props, and lighting.
- Competitive SWOT: based on product specs and the furniture market. Format each quadrant header as **Strengths**, **Weaknesses**, **Opportunities**, **Threats** (bold, on its own line). 3-4 bullet points per quadrant.

Always ground content in actual product specs. Do not invent features."""


def build_user_prompt(sku: str, product_summary: str) -> str:
    return f"""Product data for SKU {sku}:

{product_summary}

Generate all nine sections below. Use exactly these ## headings so they can be parsed correctly.

## 1. PDP Copy

## 2. SEO Keywords

## 3. Organic Social Post

## 4. Outbound Email

## 5. Headlines

## 6. Paid Social Ad

## 7. Video Script

## 8. Creative Brief

## 9. Competitive SWOT"""


def call_claude(sku: str, product_summary: str, client, brand_voice: str) -> str:
    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4000,
        system=build_system_prompt(brand_voice),
        messages=[{"role": "user", "content": build_user_prompt(sku, product_summary)}],
    )
    return response.content[0].text


def parse_sections(content: str) -> dict[str, str]:
    """Split Claude's output into a dict keyed by section name."""
    sections: dict[str, str] = {}
    current_key = None
    lines: list[str] = []

    for line in content.splitlines():
        if line.startswith("## "):
            if current_key is not None:
                sections[current_key] = "\n".join(lines).strip()
            raw = re.sub(r"^##\s+\d+\.\s*", "", line).strip()
            # Match to canonical key
            current_key = next((k for k in SECTION_KEYS if k.lower() in raw.lower()), raw)
            lines = []
        else:
            lines.append(line)

    if current_key is not None:
        sections[current_key] = "\n".join(lines).strip()

    return sections


# ── Word doc export ───────────────────────────────────────────────────────────

def _add_markdown_paragraph(doc: Document, text: str, skip_blank: bool = False) -> bool:
    """
    Write a markdown-formatted line into the doc.
    Returns True if a bullet was written (so the caller can skip blank lines between bullets).
    skip_blank: if True, blank lines are skipped entirely (no extra paragraph added).
    """
    stripped = text.strip()

    # Blank line
    if not stripped:
        if not skip_blank:
            doc.add_paragraph()
        return False

    # Horizontal rule
    if stripped in ("---", "___", "***"):
        doc.add_paragraph()
        return False

    # Bullet: "- text" or "• text"
    if re.match(r"^[-•]\s+", stripped):
        bullet_text = re.sub(r"^[-•]\s+", "", stripped)
        p = doc.add_paragraph(style="List Bullet")
        bold_match = re.match(r"\*\*(.+?)\*\*[:\s]*(.*)", bullet_text)
        if bold_match:
            p.add_run(bold_match.group(1) + ": ").bold = True
            p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", bold_match.group(2)))
        else:
            p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", bullet_text))
        return True

    # Bold-only line: **text** or **text:** — handles SWOT quadrant headers and lead sentences
    bold_match = re.match(r"^\*\*(.+?)\*\*:?\s*$", stripped)
    if bold_match:
        p = doc.add_paragraph()
        run = p.add_run(bold_match.group(1))
        run.bold = True
        p.paragraph_format.space_before = Pt(4)
        return False

    # Hashtag line
    if re.match(r"^#\w", stripped):
        p = doc.add_paragraph(stripped)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1D, 0x9B, 0xF0)
        return False

    # Script cues [VISUAL:] or [VO:]
    if re.match(r"^\[(?:VISUAL|VO)[:\s]", stripped, re.IGNORECASE):
        p = doc.add_paragraph()
        p.add_run(stripped).italic = True
        return False

    # Detect un-bolded concept/quadrant headers Claude occasionally emits
    # e.g. "Concept 2: The Bright Farmhouse Office" or "Weaknesses"
    if re.match(r"^(Concept \d+|Strengths|Weaknesses|Opportunities|Threats)[:\s]", stripped):
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        run.bold = True
        p.paragraph_format.space_before = Pt(4)
        return False

    # Plain paragraph — strip remaining markdown bold markers
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
    doc.add_paragraph(clean)
    return False


def save_docx(dest, product_name: str, sku: str, product_summary: str, sections: dict[str, str]):
    """
    Write a formatted Word document.
    dest: a file path (str/Path) or a BytesIO buffer.
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"{product_name}", level=1)
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Meta row
    meta = doc.add_paragraph()
    meta.add_run("SKU: ").bold = True
    meta.add_run(sku + "     ")
    meta.add_run("Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%B %d, %Y %I:%M %p"))

    doc.add_paragraph()

    # Product data
    doc.add_heading("Product Data", level=2)
    p = doc.add_paragraph(product_summary)
    if p.runs:
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Content sections
    doc.add_heading("Generated Content", level=2)

    for key in SECTION_KEYS:
        text = sections.get(key, "")
        if not text:
            continue

        doc.add_heading(key, level=3)

        last_was_bullet = False
        for line in text.splitlines():
            last_was_bullet = _add_markdown_paragraph(
                doc, line, skip_blank=last_was_bullet
            )

        doc.add_paragraph()

    if isinstance(dest, (str, Path)):
        doc.save(dest)
    else:
        doc.save(dest)
